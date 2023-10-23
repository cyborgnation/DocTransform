import os
from docx import Document

# specify the directory where your Word files are located
directory = './docs/'


for dirpath, dirnames, filenames in os.walk('./docs/Editable Files'):
    parent_folder_name = os.path.basename(dirpath)
    
    if parent_folder_name.startswith(("Policies", "Procedures", "Overview")):
        
        for filename in filenames:
            if filename.endswith('.docx'):
                file_path = os.path.join(dirpath, filename)
            
            # Check if file exists (although it should based on os.walk's output)
            if not os.path.exists(file_path):
                print(f"File not found: {file_path}")
                continue

            try:
                # open the document
                doc = Document(file_path)

                # Check if there's at least one table and the first table has at least one cell
                if doc.tables and doc.tables[0].rows and doc.tables[0].cell(0, 0):
                    title = doc.tables[0].cell(0, 0).text

                    # Rest of your processing ...

                else:
                    print(f"Skipped file due to missing or incomplete table: {filename}")
            except Exception as e:
                print(f"Error processing {filename}: {e}")
                
            # convert the title to title case
            title = title.title()

            # Reduce the title length
            max_title_length = 100
            title = title[:max_title_length]

            # Get the last 10 characters of the original filename
            last_10_chars = filename[-15:]

            # Remove forbidden characters
            forbidden_chars = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
            for char in forbidden_chars:
                title = title.replace(char, '')

            # Create the new filename
            new_filename = f"{title} {last_10_chars}"
                        
            # construct the new file path
            new_file_path = os.path.join(dirpath, new_filename)

            # Handle file existence
            if os.path.exists(new_file_path):
                new_file_path = new_file_path.replace('.docx', '_duplicate.docx')
                        
            # rename the file
            os.rename(file_path, new_file_path)
    elif parent_folder_name.startswith(("Exhibits", "Guide Cards", "Temporary Directives", "Memos")):
        for filename in filenames:
            file_path = os.path.join(dirpath, filename)

            # Split the filename and the file extension
            base_name, file_extension = os.path.splitext(filename)
            
            # Get all but the last 10 characters of the base name and replace underscores with spaces
            prefix = base_name[:-10].replace('_', ' ')
            
            # Get the last 10 characters of the base name
            suffix = base_name[-10:]
            
            # Construct the new filename
            new_filename = prefix + suffix + file_extension
            
            # Construct the new file path
            new_file_path = os.path.join(dirpath, new_filename)
            
            print(f"Old Path: {file_path}")
            print(f"New Path: {new_file_path}")
            
            # Rename the file
            os.rename(file_path, new_file_path)